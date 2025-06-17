from docx.api import Document
from googletrans import Translator

# Step 1: read the text from the docx file
document = Document("to be translated.docx")
# extract all text from the doc
all_text = "\n".join([p.text for p in document.paragraphs])
print(all_text)

# Step 2: translate the extracted text
translator = Translator()
target_language = 'en'
try:
    translated_text = translator.translate(all_text, dest=target_language).text
except Exception as e:
    print("Error occurred during translation:", e)
    exit()

# Step 3: output the translated text into a docx file
new_document = Document()
for paragraph in translated_text.split("\n"):
    new_document.add_paragraph(paragraph)  # add each paragraph as a new paragraph in the doc

new_document.save("translated2.docx")

print(f"Translation complete! The translated document is saved as 'translated.docx'.")