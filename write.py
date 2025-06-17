from docx import Document

document = Document()

document.add_heading("hello world!",0)

p=document.add_paragraph("sample text")

p.add_run("this is bold").bold=True
p.add_run("this is italic").italic=True

document.add_paragraph("item 1", style="List Bullet")
document.add_paragraph("item 2", style="List Bullet")


document.save("test.docx")