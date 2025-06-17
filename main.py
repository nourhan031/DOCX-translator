from docx.api import Document
from googletrans import Translator

# Step 1: Read the text from the docx file
document = Document("to be translated.docx")
all_text = "\n".join([p.text for p in document.paragraphs])
print(all_text)

# Split text into manageable chunks for translation if it's too large
chunk_size = 5000  # Number of characters per chunk
chunks = [all_text[i:i + chunk_size] for i in range(0, len(all_text), chunk_size)]

# Step 2: Translate the extracted text
translator = Translator()
target_language = 'en'
translated_text = []

try:
    for chunk in chunks:
        # Translate each chunk and append the translated text
        translated_chunk = translator.translate(chunk, dest=target_language).text
        translated_text.append(translated_chunk)
except Exception as e:
    print("Error occurred during translation:", e)
    exit()

# Join the translated chunks back into a single string
final_translated_text = "\n".join(translated_text)

# Step 3: Output the translated text into a new docx file
new_document = Document()
for paragraph in final_translated_text.split("\n"):
    new_document.add_paragraph(paragraph)  # Add each paragraph as a new paragraph in the doc

new_document.save("translated.docx")

print(f"Translation complete! The translated document is saved as 'translated.docx'.")
