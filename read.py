from docx.api import Document

document = Document("test.docx")

# for p in document.paragraphs:
#     if p.style.name.startswith("Heading") or p.style.name=="Title":
#         print(p.text)

# prints all text
all_text = ""

for p in document.paragraphs:
    all_text += p.text
    all_text += "\n"

print(all_text)